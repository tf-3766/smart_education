from __future__ import annotations

from pathlib import Path
import re
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_MD = ROOT / "在线教育辅助教学系统_项目设计文档_第一组.md"
OUT_DOCX = ROOT / "在线教育辅助教学系统_项目设计文档_第一组.docx"


COVER_MD = """# 在线教育辅助教学系统

# 项目设计文档

文档类型：概要设计与详细设计结合版

项目选题：选题 8

小组名称/组号：第一组

文档版本：V1.0

生成日期：2026-07-08

适用范围：本设计文档承接需求分析说明书，面向前端、后端业务、后端 AI、测试联调与答辩说明使用。
"""


BODY_MD = r"""
# 第一章 引言

## 1.1 编写目的

本《项目设计文档》用于承接《在线教育辅助教学系统需求分析说明书》，把已经明确的业务需求转换为可开发、可联调、可测试的系统设计方案。文档重点说明系统总体架构、服务边界、功能模块、数据库、接口、前端、安全、部署、测试和协作计划，作为第一组成员分工、编码实现、接口评审、联调验收和答辩说明的依据。

本项目设计的核心不是单纯展示课程信息，而是让学生、教师、管理员围绕课程学习、作业批改、考试安排、成绩反馈、论坛互动、公告治理和学习预警形成完整教学闭环。AI 能力只作为嵌入业务流程的辅助能力，负责生成回答、草稿、解释、引用和建议，不直接替代教师、管理员或学生完成关键业务决策。

## 1.2 项目概述

在线教育辅助教学系统面向高校教学场景，服务对象包括学生、教师和管理员。系统通过前后端分离的 B/S 架构，为学生提供选课、章节学习、资料查看、作业提交、考试安排、成绩查询、课程论坛和学习预警；为教师提供课程建设、章节课时维护、作业发布、批改评分、题库与组卷、学情跟踪和课程公告；为管理员提供用户治理、课程分类、课程审核、公告管理、论坛治理、统计分析和 AI 运行状态概览。

在智能化能力方面，系统基于课程资料和章节内容建设课程知识库，提供课程 RAG 智能答疑、章节知识点摘要、作业批改评语草稿、学习风险解释和智能组卷建议。AI 功能必须继承当前用户的业务权限和资源范围，并显示引用来源、上下文和异常状态。教师或管理员确认之前，AI 结果不得写入正式成绩、正式评语、正式课程内容、正式试卷或正式预警处理状态。

## 1.3 设计原则

1. 前后端分离：前端只负责展示、交互和状态引导，后端负责最终权限校验、业务规则和数据一致性。
2. 轻量微服务：只部署 `edu-gateway`、`edu-biz-service`、`edu-ai-service` 三个应用服务，`edu-common` 作为公共 Maven 模块，不作为独立微服务。
3. 业务与 AI 解耦：传统业务事实归 `edu-biz-service` 管理，AI 服务不直接访问业务数据库，不直接写正式业务数据。
4. 角色权限与资源权限结合：先校验角色和功能权限，再校验课程归属、选课关系、作业归属、成绩归属、资料范围和题库范围。
5. 数据一致性优先：课程、选课、作业、成绩、考试和预警等关键操作在业务服务中同步完成，必要时使用事务、乐观锁和幂等键。
6. AI 结果人工确认：AI 输出是草稿、建议、解释或引用，关键业务数据必须由教师或管理员确认后写入。
7. 接口契约优先：REST API、SSE 事件、错误码、DTO/VO、内部上下文契约先评审再开发。
8. 可维护、可测试、可扩展：模块按业务域组织，公共能力保持稳定，测试覆盖正常、异常、越权、状态冲突和 AI 降级场景。

## 1.4 参考资料

本文使用的资料以本机和当前仓库可确认内容为准，不伪造论文或不存在的资料。

- `E:\武汉理工实训选题.docx`：老师提供的实训选题说明，确认选题 8、业务范围、AI 范围和技术栈参考。
- `在线教育辅助教学系统_需求分析说明书_第1组.md` 与 `第1组-在线教育辅助教学系统需求分析说明书.docx`：确认项目名称、组号、需求边界和功能优先级。
- `design-research.md`：前端视觉调研与信息架构，确认“校园学习工作台”设计方向。
- `sitemap.md`：学生端、教师端、管理员端路由规划和 AI 业务入口。
- `ui-spec.md`：Vue 3、Element Plus、三端页面布局、AI 组件和状态展示规范。
- `wireframes.md`：学生学习首页、章节学习、教师工作台、批改工作台、管理员看板、智能组卷、学习预警等文字版低保真原型。
- `docs/backend-architecture.md`、`docs/backend-conventions.md`、`docs/database-conventions.md`、`docs/api-style.md`：后端架构、编码、数据库和 API 风格规范。
- `docs/module-ownership.md`、`docs/team-development-workflow.md`、`docs/backend-two-person-collaboration-plan.md`、`docs/mvp-scope.md`、`docs/migration-register.md`：双人后端协作、分支、迁移、MVP 范围和 owner 规则。
- `docs/course-api-contract.md`、`docs/course-module-delivery.md`、`docs/openapi/course-module.openapi.yaml`：课程与学习模块的已实现契约、交付说明和 OpenAPI。
- `backend/README.md`、`backend/pom.xml`、各服务 `application.yml`、Flyway 迁移脚本和现有 Java/Vue 代码：确认实际技术栈、服务端口、表结构、接口路径和实现状态。
- Vue、Vite、Element Plus、Spring Boot、Spring Cloud Gateway、Spring Cloud Alibaba、MyBatis-Plus、Spring AI、Qdrant 或 Milvus 官方文档：作为技术实现参考。

# 第二章 系统总体设计

## 2.1 系统设计目标

业务目标：系统要支持高校课程教学的真实闭环，覆盖用户权限、课程建设、选课学习、章节课时、资料管理、作业提交、教师批改、考试安排、成绩统计、论坛公告和学习预警，避免只形成静态课程展示页面。

教学目标：学生能够清楚知道下一步学习任务、待交作业、考试安排、成绩反馈和风险原因；教师能够连续处理课程内容、作业批改、成绩发布、学情预警和组卷；管理员能够维护基础数据、审核课程、治理内容并掌握系统运行状态。

AI 辅助目标：AI 必须嵌入课程、章节、作业、预警和组卷场景，提供有上下文、有来源、有权限边界的答疑、摘要、评语草稿、风险解释和组卷建议。AI 不直接决定成绩、发布试卷、修改课程正文或关闭预警。

协作开发目标：前后端和双人后端成员以契约先行为原则协作。后端 A 维护 Biz 主链和正式业务事实，后端 B 维护 AI、考试题库、网关和部署联调，公共模块、Flyway、Gateway、JWT 和 API 契约通过评审后合并。

## 2.2 系统总体架构

系统采用 B/S 架构和前后端分离模式。前端使用 Vue 3 + Vite + TypeScript + Element Plus，按学生端、教师端、管理员端组织路由和页面。外部请求统一进入 `edu-gateway`，网关负责路由、JWT 初步鉴权、跨域、traceId、AI 接口限流和 SSE 转发。传统业务由 `edu-biz-service` 负责，数据存储在 MySQL 8.0，并使用 Redis、RabbitMQ、Flyway 和 MyBatis-Plus。AI 能力由 `edu-ai-service` 负责，使用 Spring AI 或 LangChain4j、Qdrant 或 Milvus、Redis 短期任务状态和 SSE 流式输出。

图 2-1 系统总体架构图

```mermaid
flowchart LR
    U[学生/教师/管理员] --> FE[Vue3 前端]
    FE --> GW[edu-gateway 网关]
    GW --> BIZ[edu-biz-service 业务服务]
    GW --> AI[edu-ai-service AI 服务]
    BIZ --> DB[(MySQL 8.0: edu_biz)]
    BIZ --> Redis[(Redis: 缓存/幂等)]
    BIZ --> MQ[(RabbitMQ: 事件)]
    AI --> VDB[(Qdrant/Milvus 向量库)]
    AI --> AIRedis[(Redis: AI 短期任务)]
    AI --> LLM[大模型服务]
    AI --> MQ
    BIZ --> AI
    AI --> BIZ
    GW --> GWRedis[(Redis: AI 限流)]
    Nacos[Nacos 注册与配置] --- GW
    Nacos --- BIZ
    Nacos --- AI
```

## 2.3 服务划分设计

表 2-1 服务职责划分表

| 服务/模块 | 职责 | 不负责什么 | 主要依赖 | 对外接口类型 | 与其他服务调用关系 | 后续扩展方向 |
|---|---|---|---|---|---|---|
| `edu-gateway` | 统一入口、显式路由、JWT 验证、CORS、traceId、AI 限流、SSE 转发、网关级错误响应 | 不编排课程、作业、成绩、考试业务；不读业务数据库；不替 Biz 判断资源归属 | Spring Cloud Gateway、Redis、Nacos、edu-common | HTTP、SSE 转发 | 外部前端只访问网关；网关转发到 Biz 或 AI | 更细粒度限流、灰度路由、网关审计、SSE 超时控制 |
| `edu-biz-service` | 用户、权限、课程、选课、章节、资料、学习记录、作业、考试、成绩、论坛、公告、预警等正式业务事实 | 不直接调用大模型等待结果；不访问向量库；不把 AI 文本直接当正式业务结果 | Spring Boot、MyBatis-Plus、MySQL、Redis、RabbitMQ、Flyway、Nacos | REST API、内部 context API、业务事件 | 对 AI 提供授权后最小上下文；接收 AI 结果后由用户确认再落库 | 作业成绩闭环、考试题库、预警规则、论坛公告、统计与审计 |
| `edu-ai-service` | RAG 答疑、章节摘要、评语草稿、风险解释、组卷建议、AI 任务状态、SSE 流式输出、向量索引 | 不提供传统业务 CRUD；不连接 Biz MySQL；不修改成绩、评语、课程、试卷、预警状态 | Spring AI 或 LangChain4j、Qdrant/Milvus、Redis、RabbitMQ、OpenFeign、Nacos | `/api/v1/ai/**`、SSE、内部 Biz context 调用 | 向 Biz 请求授权上下文；通过 MQ 接收索引事件；向前端流式返回 | 真实模型接入、向量检索、提示词版本、内容安全和 AI 管理 |
| `edu-common` | 统一响应、分页、错误码接口、JWT 工具、traceId、基础异常、少量通用工具 | 不放业务 Entity、Mapper、Service、DTO 和业务状态枚举 | Maven 公共模块 | 代码依赖，不部署 | Gateway、Biz、AI 共同依赖 | 稳定技术契约、通用错误和上下文工具 |

### 2.3.1 `edu-gateway` 网关服务

`edu-gateway` 负责所有外部 HTTP 和 SSE 请求入口。它校验 JWT 的签名和有效期，清理客户端伪造的内部请求头，生成或透传 `X-Trace-Id`，按白名单将 `/api/v1/auth/**`、`/api/v1/student/**`、`/api/v1/teacher/**`、`/api/v1/admin/**` 转发到 Biz，将 `/api/v1/ai/**` 转发到 AI，并对 AI 接口执行 Redis 限流和并发控制。

网关不负责课程、作业、成绩、考试等业务决策，不访问业务数据库，也不作为唯一安全边界。下游 Biz 和 AI 服务仍需重新校验用户身份、角色、资源范围和对象状态。

### 2.3.2 `edu-biz-service` 业务服务

`edu-biz-service` 是正式业务事实来源。当前代码已经实现认证权限、课程、教师关系、选课、章节、课时、课程资料、学习记录、课程审核等基础能力。后续作业、考试、成绩、论坛、公告、预警和 Biz 侧 AI 采用记录也应在该服务内维护。

业务服务对外提供 RESTful API，对内提供 `/_internal/v1/**` 授权上下文契约给 AI 服务。每次写操作必须在服务端校验功能权限、资源归属、对象状态、版本号和幂等条件，并在必要时写入审计记录。

### 2.3.3 `edu-ai-service` AI 服务

`edu-ai-service` 是无状态或弱状态的 AI 计算服务。当前仓库已有可启动骨架、安全过滤器、traceId 和异常处理，RAG、SSE、真实模型和向量库尚未完成。设计阶段要求 AI 服务只保存自身索引、短期任务状态和运行指标，不持有 Biz 数据库账号，也不复制 Biz Entity。

AI 请求必须先经过 Gateway 限流，再由 AI 服务向 Biz 请求最小授权上下文。AI 返回 `answer`、`draft`、`suggestions`、`citations`、`warnings` 或 `taskStatus`。正式采用、保存、发布、撤回和关闭动作全部回到 Biz 服务完成。

### 2.3.4 `edu-common` 公共模块

`edu-common` 只放稳定通用能力，例如 `ApiResponse`、`PageResponse`、`ApiError`、错误码接口、JWT 签发与解析、traceId 工具和基础异常。它不能承载课程、作业、考试、AI 会话等业务模型，否则会造成 Gateway、Biz、AI 三方对业务事实的共享和误用。

## 2.4 系统核心业务流程设计

表 2-3 核心业务流程说明表

| 流程 | 参与角色 | 触发条件 | 处理步骤 | 关键权限校验 | 正常结果 | 异常情况 |
|---|---|---|---|---|---|---|
| 登录与角色进入系统 | 学生、教师、管理员 | 用户提交账号密码 | Biz 校验账号状态与 BCrypt 密码；加载角色权限；签发 JWT；前端按角色进入 `/student`、`/teacher` 或 `/admin` | 用户启用、密码正确、角色有效 | 返回 accessToken、当前用户、角色和权限 | 密码错误返回 401；无角色返回 403；禁用账号拒绝登录 |
| 教师建课、维护章节、提交审核、发布课程 | 教师、管理员 | 教师创建或修改课程 | 教师创建课程草稿；维护章节、课时和资料；负责人提交审核；管理员批准；负责人发布课程 | 教师必须是负责人；协作者不能发布；管理员审核 | 课程进入已发布或进行中状态，学生可在开放范围内选课 | 课程编号重复、版本冲突、审核未通过、非负责人越权 |
| 学生选课、学习章节、查看资料、完成课时 | 学生 | 学生浏览可选课程或进入已选课程 | 学生选课；查看已发布章节和课时；读取授权资料；开始或完成课时；系统聚合进度 | 学生本人、课程已发布、选课有效、章节课时已发布且解锁 | 形成选课记录、学习记录和进度百分比 | 未选课、资料不可见、课时未解锁、课程下线 |
| 作业发布、提交、批改、评分与反馈 | 教师、学生 | 教师发布作业，学生提交 | 教师创建作业并发布；学生保存草稿或正式提交；教师查看提交、评分、写评语并发布反馈 | 教师负责课程；学生已选课程；提交在截止和重交规则内 | 学生看到已发布成绩和评语 | 截止后提交、重复提交、越权批改、分数版本冲突 |
| 考试安排、试卷编排、考试记录和成绩发布 | 教师、学生 | 教师创建考试计划 | 教师维护题库；安排考试；编排试卷；学生按时间进入考试；交卷后教师阅卷并发布结果 | 教师负责课程；学生已选；考试时间窗有效；试卷已确认 | 学生查看考试安排和已发布结果 | 题库不足、未到开考时间、重复进入考试、成绩未发布 |
| 学习预警生成与处理 | 学生、教师、系统 | 定时规则或成绩/进度事件触发 | Biz 根据进度落后、缺交、低分生成预警；学生查看依据；教师处理或关闭 | 学生只能看本人；教师只能看负责课程；AI 仅解释授权数据 | 预警有等级、证据、处理状态和历史记录 | 数据不足、误报纠正、AI 解释失败、教师越权查看 |
| AI 课程答疑 | 学生、教师、AI 服务 | 用户在章节学习或 AI 页提问 | Gateway 限流；AI 向 Biz 获取授权课程资料上下文；检索向量库；SSE 返回回答与引用 | 用户有课程或章节访问权；资料已发布且可见 | 返回 meta、delta、citation、done | 无可靠来源、权限不足、AI 超时、SSE 中断 |
| AI 批改评语草稿 | 教师、AI 服务 | 教师在批改工作台点击生成 | Biz 提供提交、量规、问题点最小上下文；AI 生成可编辑评语草稿；教师编辑并确认 | 教师负责课程和作业；只能读取当前提交 | 生成草稿，教师确认后 Biz 写正式评语 | AI 失败不影响评分；不得自动填分或覆盖教师评语 |
| AI 章节摘要 | 教师、AI 服务、学生 | 教师在课时编辑页生成摘要 | AI 基于课程资料和课时正文生成摘要草稿；教师编辑并发布；学生只看已发布摘要 | 教师负责课程；学生仅访问已发布内容 | 摘要成为课程内容的一部分并留来源 | 资料不足、草稿未确认、章节未发布 |
| AI 智能组卷建议 | 教师、AI 服务 | 教师进入具体考试试卷编排页 | 教师设置知识点、题型、难度、分值；AI 基于授权题库给候选题和分布建议；教师采纳或修改 | 教师有题库和课程权限；候选题必须来自授权题库 | 保存教师确认后的试卷草稿 | 题库不足、不满足分值、AI 不得虚构题目或自动发布 |

图 2-2 核心业务流程图

```mermaid
flowchart TD
    A[管理员维护用户与课程分类] --> B[教师创建课程草稿]
    B --> C[维护章节/课时/资料]
    C --> D[提交课程审核]
    D --> E{管理员审核}
    E -->|通过| F[教师发布课程]
    E -->|驳回| B
    F --> G[学生选课]
    G --> H[学生学习课时/查看资料]
    H --> I[教师发布作业或考试]
    I --> J[学生提交作业/参加考试]
    J --> K[教师批改评分]
    K --> L[成绩发布与反馈]
    L --> M[学习进度与预警]
    M --> N[学生改进/教师干预]
```

图 2-3 AI 辅助教学流程图

```mermaid
sequenceDiagram
    participant UI as 前端业务页面
    participant GW as edu-gateway
    participant AI as edu-ai-service
    participant BIZ as edu-biz-service
    participant VDB as 向量库
    participant LLM as 大模型

    UI->>GW: 发起 AI 请求(课程/章节/提交/预警/考试上下文)
    GW->>GW: JWT 校验、traceId、AI 限流
    GW->>AI: 转发请求或建立 SSE
    AI->>BIZ: 请求授权后的最小上下文
    BIZ-->>AI: 上下文 DTO 或拒绝
    AI->>VDB: 检索课程资料/题库引用
    AI->>LLM: 生成回答、草稿或建议
    AI-->>UI: meta/delta/citation/done 或 error
    UI->>GW: 用户编辑后确认保存
    GW->>BIZ: 正式业务命令
    BIZ-->>UI: 保存后的正式业务结果
```

## 2.5 角色与权限总体设计

表 2-2 角色权限概览表

| 权限维度 | 学生 | 教师 | 管理员 | AI 服务权限继承规则 | 前端控制 | 后端资源级权限 |
|---|---|---|---|---|---|---|
| 系统入口 | `/student` 学习首页、课程、任务、成绩、论坛、AI 学习助手 | `/teacher` 教学工作台、课程、作业批改、考试题库、预警、互动 | `/admin` 数据看板、用户、课程治理、公告、论坛治理、统计、AI 管理 | 继承触发页面的当前用户、角色和业务上下文 | 路由守卫隐藏不可用菜单 | Biz/AI 重新校验 JWT、角色和权限 |
| 课程访问 | 只能看已选且已发布课程；可选课程只显示公开字段 | 只能管理本人负责或协作课程 | 可审核和治理课程，不默认替教师改正文 | 只能使用用户有权访问的课程资料 | 课程卡片与按钮按角色展示 | 校验选课关系、课程教师关系、课程状态 |
| 作业成绩 | 只能提交本人作业、查看本人已发布成绩 | 只能批改负责课程的提交，发布成绩 | 查看聚合统计和治理信息，不默认改分 | 可生成评语草稿，不写正式分数 | 操作按钮禁用或隐藏 | 校验作业归属、提交归属、成绩发布状态 |
| 考试题库 | 只能参加本人有资格的考试，查看已发布结果 | 管理负责课程考试与有权题库 | 查看运行监控和治理指标 | 组卷建议只使用授权题库 | 路由限制考试入口 | 校验考试时间窗、试卷状态、题库范围 |
| 论坛公告 | 学生在已选课程发帖评论，查看目标范围公告 | 教师管理本课程论坛和公告 | 管理员全局公告和内容治理 | 可解释异常或建议，不替代治理决定 | 菜单和操作区按权限显示 | 校验课程成员、发布范围、治理权限 |
| 预警与 AI | 只看本人预警和建议 | 看负责课程学生预警并记录干预 | 看统计和运行状态，不默认看私人 AI 正文 | AI 不扩大数据范围，不展示模型内部推理 | AI 面板显示上下文和来源 | 校验本人/负责课程/资料可见范围 |

前端权限控制只是体验优化，不能构成安全边界。用户手动修改 URL、请求体或资源 ID 时，后端仍必须校验课程归属、选课关系、作业归属、成绩归属、资料可见范围、题库所属课程和资源当前状态。

# 第三章 功能模块设计

第三章按业务域拆分模块。每个模块都围绕职责、角色、输入输出、规则、数据表、接口、权限、异常和模块关系展开。当前仓库已实现认证与课程学习基础能力；作业、考试、成绩、论坛、公告、预警和 AI 业务为本设计文档基于需求和项目材料补足的后续实现范围。

表 3-1 功能模块设计表

| 模块 | 模块职责 | 主要角色 | 输入输出 | 涉及数据表 | 关键接口 | 权限控制与异常处理 | 与其他模块关系 |
|---|---|---|---|---|---|---|---|
| 用户与权限 | 登录、JWT、当前用户、角色权限、资源级权限、操作审计 | 学生、教师、管理员 | 输入账号密码和 Token；输出当前用户、角色、权限、错误码 | `sys_user`、`sys_role`、`sys_permission`、`sys_user_role`、`sys_role_permission` | `/api/v1/auth/login`、`/api/v1/auth/me`、`/api/v1/auth/logout` | BCrypt、JWT、401/403、traceId、审计 | 所有模块的身份入口 |
| 课程与学习 | 课程创建、审核、发布、协作者、选课、章节、课时、资料、进度 | 学生、教师、管理员 | 输入课程、章节、课时、资料和学习命令；输出课程详情、大纲、进度 | `edu_course`、`edu_course_teacher`、`edu_course_enrollment`、`edu_course_chapter`、`edu_course_lesson`、`edu_course_material`、`edu_lesson_learning_record`、`edu_course_review` | 教师课程、学生课程、管理员审核、学习进度接口 | 课程 owner/协作者、学生选课、发布/解锁状态 | 作业、考试、AI、论坛、公告、预警的上游上下文 |
| 作业、批改与成绩 | 作业发布、提交、逾期、批改、评分、评语、成绩统计 | 学生、教师 | 输入作业、提交、评分和发布命令；输出提交状态、分数、反馈 | `edu_assignment`、`edu_assignment_submission`、`edu_grade` | `/teacher/courses/{courseId}/assignments`、`/student/assignments/{assignmentId}` | 截止时间、重交、版本冲突、教师负责课程 | 调用 AI 评语草稿，触发预警和统计 |
| 考试与题库 | 题库、考试安排、试卷编排、考试记录、成绩统计 | 学生、教师 | 输入题目、考试计划、试卷约束；输出考试安排、试卷草稿、结果 | `edu_exam`、`edu_question`、`edu_exam_paper` | `/teacher/exams`、`/teacher/question-bank`、`/student/exams` | 时间窗、题库范围、试卷发布状态 | AI 智能组卷建议，成绩模块后续接入 |
| 论坛与公告 | 课程论坛、帖子、评论、课程公告、系统公告、内容治理 | 学生、教师、管理员 | 输入帖子、评论、公告、治理命令；输出列表和状态 | `edu_forum_post`、`edu_forum_comment`、`edu_announcement` | `/student/forums`、`/teacher/forums`、`/admin/forum-moderation`、`/admin/announcements` | 课程成员可见、作者编辑、教师课程治理、管理员全局治理 | 与课程、通知、管理员治理关联 |
| 学习预警 | 规则、等级、触发依据、学生展示、教师处理、AI 解释 | 学生、教师、系统 | 输入进度、缺交、低分和处理命令；输出预警、证据、建议 | `edu_learning_warning` | `/student/progress/warnings/{warningId}`、`/teacher/warnings/{warningId}` | 只看本人或负责课程；AI 不能关闭预警 | 依赖课程、作业、成绩，调用 AI 风险解释 |
| AI 智能辅助 | 知识库构建、RAG、摘要、评语草稿、风险解释、组卷建议、引用和降级 | 学生、教师、管理员、AI 服务 | 输入业务上下文和问题；输出回答、草稿、建议、引用、任务状态 | `ai_conversation`、`ai_message`、`ai_citation`、`ai_task`；向量库集合 | `/api/v1/ai/course-qa/streams`、AI 摘要、评语、预警、组卷接口 | 继承业务权限、限流、SSE、无来源、超时、人工确认 | 嵌入课程、作业、预警、考试，不独立替代业务 |
| 管理员治理与统计 | 用户、课程分类、审核、公告、论坛治理、数据统计、AI 状态 | 管理员 | 输入治理命令和筛选；输出统计、待办、审计和状态 | 上述治理表和聚合视图 | `/admin/users`、`/admin/course-reviews`、`/admin/statistics`、`/admin/ai-management` | 管理员权限细分，敏感正文默认不展示 | 连接所有业务域的治理和运行视角 |

## 3.1 用户与权限模块设计

用户与权限模块负责登录认证、JWT 签发与解析、当前用户信息、角色与权限、资源级权限和操作日志。当前代码已使用 BCrypt 校验密码、JWT 作为无状态登录凭证，并通过 `STUDENT`、`TEACHER`、`ADMIN` 三类角色控制入口。

主要功能包括：用户登录、退出、查询当前用户、角色权限加载、方法级权限校验、资源级权限校验和审计记录。登录输入为用户名和密码，输出为 accessToken、tokenType、expiresAt、user、roles 和 permissions。异常场景包括密码错误、Token 过期、账号禁用、无角色、角色越权和资源越权。

核心规则：密码只保存 BCrypt hash；Token 有效期由服务端配置；前端不得自行构造角色；每个业务模块不能只依赖角色，还必须校验资源归属。操作日志不得记录明文密码、Token、模型 key、考试答案、学生作业全文或私人 AI 会话正文。

## 3.2 课程与学习模块设计

课程与学习模块是当前仓库已实现最完整的业务域，负责课程创建、课程审核、发布下线、教师负责人和协作者、学生选课退选、章节课时、资料元数据、学习记录和进度计算。

教师创建课程时，系统建立负责人关系并校验课程编号唯一。协作者可以维护教学内容，但不能管理课程成员、提交审核、发布或下线课程。管理员审核课程时记录审核结论、原因、审核人和时间。学生只能选择已通过审核、已发布或进行中且开放选课的课程；学习时只能访问已选课程内已发布、已解锁的章节、课时和资料。

输入包括课程信息、章节信息、课时正文、资料元数据、选课命令、学习开始和完成命令。输出包括教师课程列表、学生课程目录、课程详情、章节课时树、资料访问对象、学习记录和课程进度。进度口径为 `completedLessons / availableLessons * 100`，其中可学习课时必须已发布、所属章节已发布且已解锁。

## 3.3 作业、批改与成绩模块设计

作业、批改与成绩模块设计目标是形成“教师发布作业 -> 学生提交 -> 教师批改 -> 学生查看成绩与反馈”的真实闭环。该模块在当前代码中尚未落表实现，但属于 MVP 必须完成范围。

主要功能包括作业草稿、作业发布、附件元数据、学生保存草稿、正式提交、逾期处理、重交规则、教师批改、评分量规、评语、成绩发布、成绩统计和学生反馈查看。AI 评语草稿只作为教师批改页面的辅助入口，教师确认后才写入正式评语。

关键规则：教师只能发布负责或协作课程的作业；学生必须已选课程；截止时间以服务端时间为准；已发布成绩学生可见，未发布成绩不能按 0 分展示；已发布成绩修改需保留更正记录；AI 不自动评分、不自动发布、不覆盖教师已有评语。

## 3.4 考试与题库模块设计

考试与题库模块负责题库管理、考试安排、试卷编排、考试记录和成绩统计。MVP 可先实现考试安排、题库、试卷草稿和智能组卷建议，不实现完整在线考试引擎；若后续实现考试会话和答题记录，需要在表设计和权限设计中单独扩展。

教师维护题目时需要标注课程、知识点、题型、难度、分值、答案和解析。题库范围分为本人私有、课程共享和学校公共。考试安排包含考试名称、课程、开始时间、结束时间、时长、规则和发布状态。试卷编排可以手工选题，也可以请求 AI 生成候选题和分布建议，但最终试卷必须由教师确认。

异常场景包括题库不足、题目无权限、考试未发布、未到开考时间、重复进入考试、成绩未发布和考试取消。学生只能查看本人可参加的考试和已发布结果。

## 3.5 论坛与公告模块设计

论坛与公告模块负责课程论坛、帖子、评论、课程公告、系统公告和内容治理。课程论坛是已选学生和课程教师之间的课程内讨论区，不是全站匿名社区。公告分为课程公告和系统公告，学生按课程选课关系、公告受众和发布状态查看。

学生可在已选课程发帖、评论、编辑或删除本人未锁定内容；教师可治理本人负责课程的帖子和评论；管理员可处理全局举报、隐藏恢复内容、发布系统公告和查看治理记录。所有治理操作应保留操作者、对象、原因和时间。

异常场景包括学生访问未选课程论坛、被锁定内容再次编辑、公告撤回后继续访问、管理员越权查看私人 AI 会话正文等。

## 3.6 学习预警模块设计

学习预警模块负责规则、等级、触发依据、学生端展示、教师端处理、AI 风险解释和改进建议。首版规则可基于课程进度落后、作业缺交、低分或考试异常生成低/中/高风险预警。

预警记录由 Biz 服务生成和保存，包含课程、学生、等级、状态、触发原因、证据摘要、创建时间和处理记录。学生可以查看本人预警、依据和改进建议；教师可以查看负责课程学生的预警并记录干预。学生“已读”不等于已解决，教师关闭预警也不能删除历史证据。

AI 仅在预警详情中解释风险原因和生成建议，不改变预警等级和状态。AI 建议转为学生计划或教师干预记录时，必须由学生或教师确认。

## 3.7 AI 智能辅助模块设计

AI 智能辅助模块单独设计，但不和业务模块割裂。AI 能力必须从课程、章节、作业、预警或考试业务页面进入，自动带入上下文、权限范围和资料范围。

### 3.7.1 课程知识库构建

课程资料发布、更新或下线后，Biz 通过事件通知 AI 服务进行资料切分、Embedding、索引更新或索引删除。向量 payload 只保存检索定位所需的最小字段，如 courseId、resourceId、resourceVersion、chunkId、locator 和 accessScope，不保存成绩、完整提交、考试答案或用户敏感信息。

### 3.7.2 RAG 智能答疑

学生或教师在章节学习页提问时，AI 服务先向 Biz 获取当前用户有权访问的课程、章节和资料范围，再检索向量库并调用大模型生成回答。回答必须返回引用来源；无可靠资料时返回明确提示，不能伪造引用。

### 3.7.3 章节知识点摘要

教师在课时编辑或章节管理页触发摘要生成。AI 返回摘要草稿、来源资料和生成时间。教师编辑并发布后，摘要才作为正式课程内容展示给学生。学生不能看到未确认草稿。

### 3.7.4 作业批改评语草稿

教师在批改工作台触发评语草稿。AI 输入只包含当前提交、作业要求、评分量规、教师选择的问题点和必要课程上下文，不包含其他学生提交。AI 返回可编辑草稿，教师确认后 Biz 写入正式评语。

### 3.7.5 学习风险解释

AI 基于 Biz 提供的预警证据、进度、缺交、低分和已发布成绩解释风险，并给出具体改进建议。AI 不使用其他学生私人数据，也不输出“必定挂科”等绝对化结论。

### 3.7.6 智能组卷建议

教师设置考试、知识点、题型、难度、分值、题库范围后，AI 返回候选题和分布建议。候选题必须来自教师有权访问的题库。题库不足时列出缺口，不虚构题目。最终试卷由教师确认和保存。

### 3.7.7 AI 结果引用与人工确认

AI 输出应带 `citations`、`groundingStatus`、`resourceVersion` 和可再次鉴权的访问定位。AI 结果被教师采纳、学生加入计划或管理员处理时，Biz 应保存“最终人工确认内容”和关联的 AI 来源，而不是只保存模型原文。

### 3.7.8 AI 异常与降级处理

AI 失败、超时、无资料、权限不足或限流时，应返回明确错误并保留用户输入和业务上下文。传统课程、作业、评分、成绩、论坛和公告功能不应因 AI 不可用而中断。AI 不展示模型原始推理过程、系统提示词、模型 key 或供应商原始错误体。

## 3.8 管理员治理与统计模块设计

管理员治理与统计模块负责用户管理、课程分类、课程审核、公告管理、论坛治理、数据统计和 AI 服务运行状态概览。管理员可以查看全局聚合数据和必要治理明细，但不默认替教师修改课程正文、作业分数、考试分数，也不默认浏览学生私人 AI 会话正文。

管理员数据看板首屏建议展示总用户、运行课程、本周活跃率、待审核或异常数量、课程状态、教学活动趋势、公告审核队列、论坛举报和 AI 索引失败提醒。统计结果要能下钻到业务列表，并按学期、院系、课程和时间范围筛选。

# 第四章 数据库设计

## 4.1 数据库设计原则

业务数据库使用 MySQL 8.0，schema 建议为 `edu_biz`。表名采用小写下划线命名，认证表以 `sys_` 开头，教学业务表以 `edu_` 开头，AI 会话与任务表以 `ai_` 开头。主键统一为 `BIGINT` 雪花 ID，Java 使用 `Long`，对外 JSON 序列化为字符串。

核心表统一包含 `id`、`created_at`、`created_by`、`updated_at`、`updated_by`、`deleted`、`version` 七类审计字段。逻辑删除使用 `deleted=0/1`，并发更新使用乐观锁 `version`。数据库迁移只允许 Flyway，已有共享迁移不可修改，后续生产迁移使用秒级时间戳版本。

项目不使用数据库物理外键强绑定，原因是三服务边界要求清晰、后续可扩展为跨服务契约，同时避免误用级联删除破坏历史学习记录、成绩、提交和审计。关联完整性通过应用服务事务、唯一约束、复合索引、状态校验和一致性审计保证。

## 4.2 核心实体关系

图 4-1 数据库核心 ER 图

```mermaid
erDiagram
    sys_user ||--o{ sys_user_role : has
    sys_role ||--o{ sys_user_role : assigned
    sys_role ||--o{ sys_role_permission : grants
    sys_permission ||--o{ sys_role_permission : contains
    sys_user ||--o{ edu_course_teacher : teaches
    edu_course ||--o{ edu_course_teacher : has
    sys_user ||--o{ edu_course_enrollment : enrolls
    edu_course ||--o{ edu_course_enrollment : selected_by
    edu_course ||--o{ edu_course_chapter : contains
    edu_course_chapter ||--o{ edu_course_lesson : contains
    edu_course ||--o{ edu_course_material : owns
    edu_course_lesson ||--o{ edu_lesson_learning_record : tracked_by
    sys_user ||--o{ edu_lesson_learning_record : learns
    edu_course ||--o{ edu_course_review : reviewed
    edu_course ||--o{ edu_assignment : has
    edu_assignment ||--o{ edu_assignment_submission : receives
    sys_user ||--o{ edu_assignment_submission : submits
    edu_assignment_submission ||--o{ edu_grade : graded_as
    edu_course ||--o{ edu_exam : schedules
    edu_course ||--o{ edu_question : owns
    edu_exam ||--o{ edu_exam_paper : has
    edu_question ||--o{ edu_exam_paper : selected
    edu_course ||--o{ edu_forum_post : discusses
    edu_forum_post ||--o{ edu_forum_comment : has
    edu_course ||--o{ edu_announcement : publishes
    edu_course ||--o{ edu_learning_warning : warns
    sys_user ||--o{ edu_learning_warning : receives
    sys_user ||--o{ ai_conversation : starts
    ai_conversation ||--o{ ai_message : contains
    ai_message ||--o{ ai_citation : cites
    ai_conversation ||--o{ ai_task : tracks
```

## 4.3 核心表设计

表 4-1 核心数据表设计表

| 表名 | 中文含义 | 主要字段 | 主键 | 关键唯一约束 | 关键索引 | 说明 |
|---|---|---|---|---|---|---|
| `sys_user` | 用户表 | username、password_hash、display_name、user_status、审计字段 | id | `uk_user_username` | `idx_user_status_deleted` | 已实现；保存学生、教师、管理员基础账号 |
| `sys_role` | 角色表 | role_code、role_name、enabled、审计字段 | id | `uk_role_code` | role_code | 已实现；内置 STUDENT、TEACHER、ADMIN |
| `sys_permission` | 权限表 | permission_code、permission_name、enabled、审计字段 | id | `uk_permission_code` | permission_code | 已实现；保存稳定功能权限码 |
| `sys_user_role` | 用户角色关系 | user_id、role_id、审计字段 | id | `uk_user_role(user_id,role_id)` | user_id、role_id | 已实现；支持用户拥有多个角色 |
| `sys_role_permission` | 角色权限关系 | role_id、permission_id、审计字段 | id | `uk_role_permission(role_id,permission_id)` | role_id、permission_id | 已实现；角色到功能权限映射 |
| `edu_course` | 课程表 | course_code、name、summary、owner_teacher_id、status、review_status、term、start_at、end_at | id | `uk_course_code` | owner/status、review、catalog、term/category | 已实现；课程运行状态与审核状态分离 |
| `edu_course_teacher` | 课程教师关系 | course_id、teacher_id、role、审计字段 | id | `uk_course_teacher(course_id,teacher_id)` | teacher、course_role | 已实现；OWNER/COLLABORATOR |
| `edu_course_enrollment` | 学生选课关系 | course_id、student_id、status、enrolled_at、withdrawn_at | id | `uk_course_enrollment(course_id,student_id)` | student_status、course_status | 已实现；退选用状态流转保留历史 |
| `edu_course_chapter` | 课程章节 | course_id、title、description、sort_order、status、published_at | id | 无全局唯一 | course_order、course_status | 已实现；章节逻辑删除 |
| `edu_course_lesson` | 课时表 | course_id、chapter_id、title、content_type、content、video_url、unlock_type、status | id | 无全局唯一 | chapter_order、course_status、unlock | 已实现；学生最小学习单元 |
| `edu_course_material` | 课程资料元数据 | course_id、chapter_id、lesson_id、name、material_type、file_key、file_url、visibility、status | id | 无全局唯一 | course、chapter、lesson | 已实现；只保存元数据，不存二进制 |
| `edu_lesson_learning_record` | 课时学习记录 | course_id、chapter_id、lesson_id、student_id、status、started_at、completed_at、study_seconds | id | `uk_lesson_student(lesson_id,student_id)` | student_course、course_lesson | 已实现；学习历史不随课时删除级联 |
| `edu_course_review` | 课程审核历史 | course_id、review_status、reviewer_id、reason、remark、reviewed_at | id | 无 | course_history、review_status | 已实现；追加式保存审核结论 |
| `edu_assignment` | 作业表 | course_id、lesson_id、title、description、deadline_at、max_score、assignment_status | id | 课程内标题可选唯一 | course_status_deadline | 设计表；教师发布作业 |
| `edu_assignment_submission` | 作业提交表 | assignment_id、student_id、submission_status、content、submitted_at、attempt_no、late_flag | id | assignment_id + student_id + attempt_no | assignment_status、student_assignment | 设计表；保存草稿、正式提交和重交 |
| `edu_grade` | 成绩表 | course_id、assignment_id、submission_id、student_id、score、comment、publication_status | id | submission_id 可唯一 | course_student、assignment_publication | 设计表；教师确认后写正式分数和评语 |
| `edu_exam` | 考试安排表 | course_id、title、start_at、end_at、duration_minutes、exam_status、rules | id | course_id + title 可选唯一 | course_status_time | 设计表；MVP 首版以安排和结果为主 |
| `edu_question` | 题目表 | course_id、owner_teacher_id、question_type、difficulty、stem、answer、analysis、scope | id | 题目编码可选唯一 | course_scope、owner_type | 设计表；题库来源和权限范围 |
| `edu_exam_paper` | 试卷/试卷题目表 | exam_id、question_id、score、sort_order、paper_status、source_type | id | exam_id + question_id + sort_order | exam_order、question | 设计表；可保存教师确认后的试卷草稿 |
| `edu_forum_post` | 论坛帖子 | course_id、author_id、title、content、post_status、pinned | id | 无 | course_status_time、author | 设计表；课程内讨论 |
| `edu_forum_comment` | 论坛评论 | post_id、course_id、author_id、content、comment_status、parent_comment_id | id | 无 | post_time、course_author | 设计表；回复、治理和审计 |
| `edu_announcement` | 公告表 | course_id、publisher_id、title、content、scope、announcement_status、published_at | id | 无 | scope_status_time、course_time | 设计表；课程公告和系统公告 |
| `edu_learning_warning` | 学习预警表 | course_id、student_id、warning_level、warning_status、reason、evidence_json、handled_by | id | course_id + student_id + active_rule 可选唯一 | student_status、course_level | 设计表；预警证据和处理状态 |
| `ai_conversation` | AI 会话表 | user_id、course_id、lesson_id、context_type、conversation_status、last_message_at | id | 无 | user_context、course_time | 设计表；Biz 保存正式会话记录 |
| `ai_message` | AI 消息表 | conversation_id、role、content、message_status、sequence、grounding_status | id | conversation_id + sequence | conversation_sequence | 设计表；保存用户问题和 AI 可展示回答 |
| `ai_citation` | AI 引用来源表 | message_id、source_id、resource_id、resource_version、title、locator_json、snippet | id | message_id + citation_no | resource_version | 设计表；展示可追溯来源 |
| `ai_task` | AI 任务表 | conversation_id、task_type、task_status、started_at、finished_at、error_code、retry_count | id | 无 | status_time、conversation | 设计表；跟踪批量摘要、组卷、索引等任务 |

## 4.4 数据状态与枚举设计

表 4-2 状态枚举设计表

| 枚举 | 取值 | 含义与主要流转 |
|---|---|---|
| 用户状态 | `ENABLED`、`DISABLED` | 启用账号可登录，禁用账号拒绝登录；管理员可启停 |
| 课程状态 | `DRAFT`、`PENDING_REVIEW`、`PUBLISHED`、`ONGOING`、`FINISHED`、`OFFLINE` | 草稿提交审核；审核通过后发布；进行中后结束；非下线状态可下线 |
| 课程审核状态 | `NOT_SUBMITTED`、`PENDING`、`APPROVED`、`REJECTED` | 未提交到审核中；管理员批准或驳回；驳回后可再提交 |
| 选课状态 | `ENROLLED`、`WITHDRAWN`、`COMPLETED` | 已选课可退选或完成；退选保留历史 |
| 章节状态 | `DRAFT`、`PUBLISHED`、`OFFLINE` | 教师草稿发布，已发布可下线 |
| 课时状态 | `DRAFT`、`PUBLISHED`、`OFFLINE` | 课时必须在章节已发布后才能供学生学习 |
| 课程资料状态 | `DRAFT`、`PUBLISHED`、`OFFLINE` | 已发布资料可被授权用户访问，下线触发索引清理 |
| 学习状态 | `NOT_STARTED`、`IN_PROGRESS`、`COMPLETED` | 开始学习后进入进行中，完成后不可退回 |
| 作业状态 | `DRAFT`、`PUBLISHED`、`CLOSED` | 教师草稿发布；截止或手动关闭后停止提交 |
| 提交状态 | `DRAFT`、`SUBMITTED`、`RETURNED`、`GRADED` | 学生草稿提交；教师可退回或批改；已批改不等于成绩已发布 |
| 成绩发布状态 | `DRAFT`、`PUBLISHED`、`WITHDRAWN` | 未发布成绩学生不可见；发布后如撤回需原因和历史 |
| 考试状态 | `DRAFT`、`PUBLISHED`、`ENDED`、`RESULTS_PUBLISHED`、`CANCELLED` | 教师发布考试；结束后发布结果；取消需原因 |
| 预警等级 | `LOW`、`MEDIUM`、`HIGH` | 由规则依据计算；AI 只解释，不直接修改等级 |
| 预警状态 | `OPEN`、`ACKNOWLEDGED`、`IN_PROGRESS`、`RESOLVED`、`CLOSED` | 学生可已读，教师可处理和关闭，证据保留 |
| AI 任务状态 | `PENDING`、`RUNNING`、`SUCCEEDED`、`FAILED`、`CANCELLED` | AI 服务短期状态或 Biz 正式记录；不代表业务结果已采纳 |

# 第五章 接口设计

## 5.1 接口设计原则

接口统一采用 RESTful API，公共前缀为 `/api/v1`，内部服务接口使用 `/_internal/v1` 且不得由 Gateway 对外暴露。响应使用统一 `ApiResponse<T>`，分页使用 `PageResponse<T>`。请求 DTO、应用层命令、响应 VO 与数据库 Entity 分离，前端不直接依赖数据库字段。

所有 Long ID 对外返回字符串，时间使用带时区 RFC 3339。查询接口统一使用 `page`、`size`、`keyword`、`status`、`sort` 等参数，排序字段由后端白名单控制。后端必须执行参数校验、角色权限、功能权限、资源范围、对象状态和 traceId 记录。错误使用正确 HTTP 状态和稳定错误码，不使用永远 HTTP 200 表示失败。

## 5.2 统一响应格式

表 5-1 统一响应格式说明表

| 场景 | HTTP | 响应结构示例 | 说明 |
|---|---:|---|---|
| 成功 | 200/201 | `{"code":"SUCCESS","message":"OK","data":{},"errors":[],"traceId":"...","timestamp":"..."}` | 创建接口使用 201，data 至少返回资源 ID |
| 分页 | 200 | `data.records/page/size/total/totalPages` | records 为空仍返回 200 |
| 参数错误 | 400 | `code=PARAM_VALIDATION_ERROR`，errors 包含 field 和 reason | 不回显密码、Token、作业全文或考试答案 |
| 未登录 | 401 | `code=UNAUTHORIZED` 或 `TOKEN_EXPIRED` | 前端跳转登录或刷新登录态 |
| 无权限 | 403 | `code=FORBIDDEN` | 不泄露其他学生、私有课程或试题是否存在 |
| 资源不存在 | 404 | `code=RESOURCE_NOT_FOUND` | 敏感资源也可按安全策略返回 404 |
| 状态冲突 | 409 | `code=RESOURCE_CONFLICT` 或 `OPERATION_NOT_ALLOWED` | 用于乐观锁、重复提交、截止后提交、状态不允许 |
| AI 限流 | 429 | `code=AI_RATE_LIMITED`，可带 `Retry-After` | Gateway 按用户、功能和时间窗限流 |
| AI 不可用 | 503 | `code=AI_SERVICE_UNAVAILABLE` | 传统业务继续可用，前端保留输入和上下文 |

## 5.3 核心接口清单

表 5-2 核心接口清单表

| 模块 | 接口名称 | 方法与路径 | 使用角色 | 请求参数 | 响应数据 | 权限说明 | 备注 |
|---|---|---|---|---|---|---|---|
| 认证与用户 | 登录 | `POST /api/v1/auth/login` | 学生、教师、管理员 | username、password | LoginVO | 账号启用、密码正确 | 当前已实现 |
| 认证与用户 | 当前用户 | `GET /api/v1/auth/me` | 已登录用户 | Bearer Token | CurrentUserVO | Token 有效 | 当前已实现 |
| 课程与学习 | 教师创建课程 | `POST /api/v1/teacher/courses` | 教师 | CreateCourseRequest | CourseDetailVO | TEACHER，当前用户成为负责人 | 当前已实现 |
| 课程与学习 | 教师课程列表 | `GET /api/v1/teacher/courses` | 教师 | page、size、keyword、status | PageResponse | 只列本人关系课程 | 当前已实现 |
| 课程与学习 | 更新课程 | `PUT /api/v1/teacher/courses/{courseId}` | 教师 | UpdateCourseRequest | CourseDetailVO | 负责人，草稿或驳回状态 | 当前已实现 |
| 课程与学习 | 提交审核 | `POST /api/v1/teacher/courses/{courseId}/submit-review` | 教师 | courseId | CourseDetailVO | 负责人，状态允许 | 当前已实现 |
| 课程与学习 | 管理员审核 | `POST /api/v1/admin/course-reviews/{courseId}/approve` | 管理员 | remark | CourseReviewVO | ADMIN，审核中 | 当前已实现 |
| 课程与学习 | 学生课程目录 | `GET /api/v1/student/courses/catalog` | 学生 | page、size、keyword | PageResponse | 已通过且可选范围 | 当前已实现 |
| 课程与学习 | 学生选课 | `POST /api/v1/student/courses/{courseId}/enroll` | 学生 | courseId | EnrollmentVO | 课程可选，学生本人 | 当前已实现 |
| 课程与学习 | 学习大纲 | `GET /api/v1/student/courses/{courseId}/outline` | 学生 | courseId | CourseOutlineVO | 已选且课程可学习 | 当前已实现 |
| 课程与学习 | 完成课时 | `POST /api/v1/student/lessons/{lessonId}/complete` | 学生 | lessonId | LearningRecordVO | 已选、已发布、已解锁 | 当前已实现 |
| 作业与成绩 | 教师发布作业 | `POST /api/v1/teacher/courses/{courseId}/assignments` | 教师 | title、deadlineAt、maxScore | AssignmentVO | 负责或协作课程 | 设计接口 |
| 作业与成绩 | 学生作业列表 | `GET /api/v1/student/assignments` | 学生 | page、courseId、status | PageResponse | 本人已选课程 | 设计接口 |
| 作业与成绩 | 学生提交作业 | `POST /api/v1/student/assignments/{assignmentId}/submissions` | 学生 | content、attachments、submitMode | SubmissionVO | 本人、未截止、允许重交 | 设计接口 |
| 作业与成绩 | 教师查看提交 | `GET /api/v1/teacher/assignments/{assignmentId}/submissions` | 教师 | page、status | PageResponse | 负责课程 | 设计接口 |
| 作业与成绩 | 教师评分发布 | `POST /api/v1/teacher/submissions/{submissionId}/grading` | 教师 | score、comment、publish | GradeVO | 负责课程、版本校验 | 设计接口 |
| 考试与题库 | 创建考试 | `POST /api/v1/teacher/courses/{courseId}/exams` | 教师 | title、startAt、endAt、rules | ExamVO | 负责课程 | 设计接口 |
| 考试与题库 | 题库管理 | `POST /api/v1/teacher/question-bank/questions` | 教师 | type、difficulty、stem、answer | QuestionVO | 题库编辑权 | 设计接口 |
| 考试与题库 | 试卷编排 | `POST /api/v1/teacher/exams/{examId}/paper` | 教师 | questions、scores | ExamPaperVO | 负责课程，考试草稿 | 设计接口 |
| 考试与题库 | 学生考试列表 | `GET /api/v1/student/exams` | 学生 | page、courseId、status | PageResponse | 本人已选课程 | 设计接口 |
| 论坛与公告 | 课程发帖 | `POST /api/v1/student/forums/posts` | 学生 | courseId、title、content | ForumPostVO | 已选课程成员 | 设计接口 |
| 论坛与公告 | 帖子评论 | `POST /api/v1/student/forums/posts/{postId}/comments` | 学生、教师 | content | CommentVO | 课程成员 | 设计接口 |
| 论坛与公告 | 发布课程公告 | `POST /api/v1/teacher/courses/{courseId}/announcements` | 教师 | title、content、scope | AnnouncementVO | 负责课程 | 设计接口 |
| 论坛与公告 | 管理员公告 | `POST /api/v1/admin/announcements` | 管理员 | title、content、audience | AnnouncementVO | 公告管理权限 | 设计接口 |
| 预警 | 学生预警详情 | `GET /api/v1/student/progress/warnings/{warningId}` | 学生 | warningId | WarningDetailVO | 本人预警 | 设计接口 |
| 预警 | 教师处理预警 | `POST /api/v1/teacher/warnings/{warningId}/actions` | 教师 | action、remark | WarningActionVO | 负责课程 | 设计接口 |
| AI 服务 | 课程答疑流 | `POST /api/v1/ai/course-qa/streams` | 学生、教师 | courseId、lessonId、question | SSE | 继承课程/章节权限 | 设计接口 |
| AI 服务 | 生成评语草稿 | `POST /api/v1/ai/grading-comments` | 教师 | submissionId、rubric、points | AiDraftVO 或 SSE | 负责课程和提交 | 设计接口 |
| AI 服务 | 章节摘要 | `POST /api/v1/ai/chapter-summaries` | 教师 | courseId、chapterId、lessonIds | AiDraftVO | 负责课程 | 设计接口 |
| AI 服务 | 风险解释 | `POST /api/v1/ai/warning-explanations` | 学生、教师 | warningId | AiSuggestionVO | 本人或负责课程 | 设计接口 |
| AI 服务 | 组卷建议 | `POST /api/v1/ai/paper-suggestions` | 教师 | examId、constraints | AiPaperSuggestionVO | 负责课程和题库 | 设计接口 |
| 管理员治理 | 用户管理 | `GET/POST /api/v1/admin/users` | 管理员 | 筛选或用户信息 | PageResponse/UserVO | 用户管理权限 | 设计接口 |
| 管理员治理 | 课程分类 | `GET/POST /api/v1/admin/course-categories` | 管理员 | name、sort、enabled | CategoryVO | 课程配置权限 | 设计接口 |
| 管理员治理 | 数据统计 | `GET /api/v1/admin/statistics` | 管理员 | term、department、timeRange | StatisticsVO | 统计权限 | 设计接口 |
| 管理员治理 | AI 管理 | `GET /api/v1/admin/ai-management/status` | 管理员 | timeRange、feature | AiOpsVO | AI 管理权限 | 设计接口 |

## 5.4 AI SSE 流式接口设计

AI 流式输出使用 `text/event-stream;charset=UTF-8`。正常顺序为 `meta -> delta* -> citation* -> done`，异常顺序为 `meta -> delta* -> citation* -> error -> close`。心跳使用 SSE comment，不定义为业务事件。

### `meta`

用途：声明 traceId、conversationId、messageId、业务上下文和初始状态。字段包括 traceId、conversationId、messageId、context、status。

```text
event: meta
id: trace-001:1
data: {"traceId":"trace-001","conversationId":"2001","messageId":"3001","context":{"courseId":"21001","lessonId":"23001","scope":"CURRENT_LESSON"},"status":"RETRIEVING"}
```

### `delta`

用途：返回面向用户的回答增量。字段包括 sequence、text。不得包含模型原始推理链、隐藏 prompt 或供应商原始事件。

```text
event: delta
id: trace-001:2
data: {"sequence":1,"text":"二叉树是一种每个节点最多有两个子节点的树结构。"}
```

### `citation`

用途：返回来源引用。字段包括 citationId、resourceId、resourceVersion、resourceType、title、locator、snippet、accessUrl。

```text
event: citation
id: trace-001:3
data: {"citation":{"citationId":"c1","resourceId":"24001","resourceVersion":3,"resourceType":"COURSE_PDF","title":"数据结构课程讲义","locator":{"kind":"PAGE","label":"第18页","page":18},"snippet":"二叉树中每个结点的度不大于2","accessUrl":"/api/v1/student/materials/24001?anchor=page-18"}}
```

### `done`

用途：声明生成结束。字段包括 finishReason、citationCount、groundingStatus、usage。usage 可用于教师或管理员视图，但不暴露供应商密钥和内部 prompt。

```text
event: done
id: trace-001:4
data: {"finishReason":"STOP","citationCount":1,"groundingStatus":"GROUNDED","usage":{"inputTokens":820,"outputTokens":196}}
```

### `error`

用途：声明流已失败并关闭连接。字段包括 code、message、retryable、partial、lastSequence、traceId。客户端保留已生成内容和用户输入。

```text
event: error
id: trace-001:4
data: {"code":"SSE_STREAM_INTERRUPTED","message":"回答未完整生成，已保留现有内容","retryable":true,"partial":true,"lastSequence":1,"traceId":"trace-001"}
```

# 第六章 前端设计

## 6.1 前端总体结构

前端设计采用 Vue 3 + Vite + TypeScript + Element Plus + Vue Router + Pinia + Axios。当前仓库 `package.json` 已安装 Vue 3、Vite、TypeScript、Element Plus、Vue Router 和 Element Plus 图标库，Pinia 与 Axios 属于后续按设计补齐的状态管理和接口封装能力。

前端按学生端、教师端、管理员端组织路由域，公共布局包含顶部栏、侧边菜单、消息入口、账户菜单、角色入口和内容区。登录态由 Token、当前用户、角色和权限共同决定。路由守卫负责体验层面的菜单隐藏和跳转，但后端仍进行最终鉴权。

Axios 封装统一注入 `Authorization`、`X-Trace-Id`，统一处理 401、403、404、409、429 和 503。Pinia 建议拆分 auth、layout、notification、course、ai 等 store，避免每个页面重复维护用户和上下文状态。

## 6.2 路由设计

图 6-1 前端路由结构图

```mermaid
flowchart TD
    Root[/] --> Login[/login]
    Root --> Student[/student]
    Root --> Teacher[/teacher]
    Root --> Admin[/admin]
    Root --> Forbidden[/403]
    Root --> NotFound[/404]
    Student --> SD[/student/dashboard]
    Student --> SC[/student/courses]
    Student --> SL[/student/courses/:courseId/lessons/:lessonId]
    Student --> SA[/student/assignments]
    Student --> SE[/student/exams]
    Student --> SG[/student/grades]
    Student --> SF[/student/forums]
    Student --> SAI[/student/ai-assistant]
    Teacher --> TD[/teacher/dashboard]
    Teacher --> TC[/teacher/courses]
    Teacher --> TG[/teacher/assignments/:assignmentId/grading/:submissionId]
    Teacher --> TP[/teacher/exams/:examId/paper]
    Teacher --> TW[/teacher/warnings]
    Admin --> AD[/admin/dashboard]
    Admin --> AU[/admin/users]
    Admin --> ACR[/admin/course-reviews]
    Admin --> AA[/admin/announcements]
    Admin --> AF[/admin/forum-moderation]
    Admin --> AS[/admin/statistics]
    Admin --> AIM[/admin/ai-management]
```

主要根路径包括 `/student`、`/teacher`、`/admin`、`/login`、`/403`、`/404`。学生端以学习首页、我的课程、学习任务、成绩与进度、互动交流和 AI 学习助手为主；教师端以教学工作台、课程管理、作业与批改、考试与题库、学情与预警、课程互动为主；管理员端以数据看板、用户管理、课程治理、内容治理、数据统计、AI 管理和系统设置为主。

## 6.3 页面设计原则

前端采用“校园学习工作台”作为统一设计方向。学生端第一屏先展示待办、继续学习、考试提醒、学习进度和风险；教师端第一屏先展示待发布、待批改、课程运行和预警；管理员端第一屏先展示全局统计、待审核、待治理和系统异常。

AI 出现在业务上下文中，而不是孤立聊天页。章节学习页有 AI 答疑侧栏；教师批改页有 AI 评语草稿；预警详情页有 AI 风险解释；试卷编排页有 AI 组卷建议。完整 AI 学习助手页只用于历史会话和延续课程上下文。

## 6.4 关键页面设计

学生学习首页：展示今日待办、继续学习、课程进度、作业考试提醒、学习风险和最新公告。AI 建议必须基于真实进度、作业和成绩，不在首页生成无上下文聊天。

学生课程详情：展示课程基本信息、进度、章节目录、近期作业、考试、公告和论坛入口。课程详情负责总览，章节学习负责沉浸式内容。

章节学习 + AI 侧栏：桌面端使用章节目录、学习内容、AI 侧栏三栏布局。AI 顶部显示当前课程、章节和资料范围，回答提供来源或无可靠来源提示。

学生作业详情：展示作业要求、截止时间、附件、保存草稿、确认提交、提交历史和教师反馈。正式提交需要二次确认。

教师教学工作台：展示负责课程、待批改、待发布、待处理预警、今日任务、课程运行概览和 AI 快捷入口。AI 快捷入口只负责定位业务，不直接生成脱离课程的内容。

教师批改工作台 + AI 评语：三栏结构为学生队列、提交内容和评分面板。AI 草稿位于评语区附近，必须可编辑，教师确认后才保存或发布。

管理员数据看板：展示总用户、运行课程、本周活跃率、待审核/异常、课程状态、教学活动趋势、公告审核、论坛举报和 AI 索引失败提醒。管理员不默认进入教师批改页改分，也不默认查看学生私人 AI 会话正文。

# 第七章 安全设计

## 7.1 身份认证安全

用户密码使用 BCrypt hash 保存。系统使用 JWT 作为无状态认证凭证，Token 包含 userId、username、activeRole、roles、permissions、issuer、expiresAt 等必要信息。Token 过期后返回 401，前端跳转登录或按后续刷新令牌机制处理。登录、登出、角色切换和异常认证都应写入安全审计。

## 7.2 角色权限控制

系统角色包括学生、教师、管理员。角色只决定入口和功能权限，不能直接代表资源归属。学生访问 `/teacher` 和 `/admin` 路由返回 403；教师访问其他教师课程时必须继续进行资源校验；管理员能力也应细分为用户管理、课程审核、内容治理、统计、系统配置和 AI 管理。

## 7.3 资源级权限控制

资源级权限包括课程负责人/协作者关系、学生选课关系、作业所属课程、提交所属学生、成绩所属学生、资料可见范围、题库所属课程、考试时间窗、公告受众和论坛课程成员。所有资源 ID 都不能只相信前端传值，后端必须重新读取当前状态。

## 7.4 数据安全

业务数据库不保存明文密码、Token、模型 key 和文件二进制。敏感数据在日志中脱敏，错误响应不暴露 SQL、堆栈、内部路径、数据库约束名和供应商错误体。成绩、考试答案、私人 AI 会话和学生提交内容按最小必要原则展示。

## 7.5 文件资料安全

文件上传成功不等于课程资料已发布或作业已提交。服务端校验文件类型、大小、用途、课程归属和访问权限。文件名只用于展示，存储使用不可预测 object key。下载和预览必须经过鉴权接口或短时签名 URL，不暴露内部路径和对象存储密钥。

## 7.6 接口安全

生产环境使用 HTTPS，CORS 使用白名单，写操作支持幂等键和乐观锁。参数校验覆盖格式、长度、跨字段、时间顺序、状态流转和服务端时间规则。AI 接口按用户、功能和时间窗口限流，SSE 配置超时、取消和断流处理。

## 7.7 AI 安全

AI 只访问当前用户有权访问的课程资料、提交内容、预警证据和题库范围。AI 不展示模型内部推理、系统提示词、密钥或未授权资料片段。AI 输出必须显示“AI 草稿”或“AI 建议”，关键结果人工确认后才进入正式业务数据。

## 7.8 日志审计与异常处理

系统记录课程审核、课程发布/下线、作业发布、成绩发布或更正、考试发布、预警处理、权限变更、论坛治理和 AI 结果采用等关键操作。审计记录包含操作人、角色、资源类型、资源 ID、动作、结果、时间和 traceId。未知异常记录服务端日志，对外返回稳定错误码。

表 7-1 安全设计措施表

| 安全项 | 设计措施 | 适用模块 | 异常处理 |
|---|---|---|---|
| 密码安全 | BCrypt hash，不记录明文密码 | 用户与权限 | 登录失败返回统一提示 |
| Token 安全 | JWT 签名、有效期、issuer、网关和服务双重校验 | Gateway、Biz、AI | 过期返回 `TOKEN_EXPIRED` |
| 角色权限 | 学生、教师、管理员入口和功能权限分离 | 全系统 | 越权返回 403 |
| 资源权限 | 课程归属、选课、作业、成绩、资料、题库逐项校验 | Biz、AI | 不可访问返回 403 或 404 |
| 数据保护 | 不暴露 SQL、堆栈、内部路径和敏感正文 | 全系统 | 返回稳定错误码 |
| 文件安全 | 权限下载、短时链接、MIME 和大小校验 | 课程资料、作业附件 | 上传失败不影响正式提交 |
| AI 限流 | Gateway 对 `/api/v1/ai/**` 限流和并发控制 | AI 服务 | 返回 429 和重试提示 |
| AI 权限 | AI 通过 Biz context 获取授权数据 | AI 服务 | 无权限、无资料、超时均可降级 |
| 日志审计 | 高风险操作写审计，日志带 traceId | 全系统 | 日志不记录密钥、Token、答案 |

# 第八章 部署与运行设计

本地开发环境使用 Windows + JDK 21 + Maven Wrapper + Node/Vite。后端依赖 MySQL 8.0、Redis、RabbitMQ 和 Nacos；当前 `deploy/docker-compose.yml` 已编排 MySQL、Redis、RabbitMQ、Nacos，尚未加入 Qdrant 或 Milvus。AI 阶段默认选择 Qdrant 作为向量库，若改用 Milvus 需新增 ADR。

启动顺序为：基础设施 MySQL/Redis/RabbitMQ/Nacos -> `edu-biz-service` -> `edu-ai-service` -> `edu-gateway` -> 前端 Vite。默认端口为 Gateway `18080`、Biz `18081`、AI `18082`、前端 `5173`。前端只访问 Gateway。

配置文件通过环境变量和 `application-local.yml`、`application-dev.yml` 管理。真实密钥、数据库密码、JWT secret、模型 key 和对象存储凭据不提交 Git。日志统一包含 traceId，便于从前端错误响应追踪到 Gateway、Biz、AI 和 MQ 消费日志。

图 8-1 部署架构图

```mermaid
flowchart LR
    Browser[浏览器] --> FE[Vite/Vue 前端]
    FE --> GW[edu-gateway:18080]
    GW --> BIZ[edu-biz-service:18081]
    GW --> AI[edu-ai-service:18082]
    BIZ --> MySQL[(MySQL 8.0 edu_biz)]
    BIZ --> Redis[(Redis)]
    BIZ --> Rabbit[(RabbitMQ)]
    AI --> Qdrant[(Qdrant/Milvus)]
    AI --> AiRedis[(Redis AI namespace)]
    AI --> LLM[大模型 Provider]
    BIZ <--> Rabbit
    AI <--> Rabbit
    Nacos[Nacos:8848] --- GW
    Nacos --- BIZ
    Nacos --- AI
```

部署环境分为开发、测试和演示。开发环境可使用 local profile 和本地种子数据；测试环境应运行完整 Flyway、接口测试、权限测试和 AI fake adapter；演示环境应冻结依赖版本、演示账号、课程、作业、成绩、预警和 AI 示例数据。数据库备份至少在执行迁移和演示前完成，AI 向量索引可按资料版本重建，不作为唯一业务事实备份。

# 第九章 测试设计

测试设计围绕功能正确性、权限安全、数据迁移、接口契约、前后端联调、AI 输出和演示稳定性展开。高风险场景包括资源越权、状态冲突、截止时间、重复提交、成绩发布、AI 无来源、SSE 中断和人工确认。

表 9-1 测试设计表

| 测试项 | 测试目标 | 测试方法 | 预期结果 |
|---|---|---|---|
| 单元测试 | 校验状态机、权限策略、进度计算和错误码 | JUnit 覆盖枚举流转、领域规则、工具类 | 正常、边界和非法状态均有明确结果 |
| 接口测试 | 校验 REST 路径、请求、响应、HTTP 状态和错误码 | MockMvc、WebTestClient、OpenAPI/Postman | 契约与实现一致 |
| 权限测试 | 防止跨角色和跨资源访问 | 用学生、教师、管理员和无关教师账号替换资源 ID | 越权返回 403 或 404 |
| 数据库迁移测试 | 保证空库和升级迁移可执行 | Flyway validate、MySQL Testcontainers、H2 快速测试 | V1-V3 和后续迁移不漂移 |
| 前后端联调测试 | 校验页面入口、路由、列表、详情和提交闭环 | 前端连接 Gateway，使用演示账号逐流程验证 | 三端核心页面可完成业务 |
| 作业闭环测试 | 覆盖发布、提交、批改、成绩发布 | 集成测试和手工验收 | 学生只看本人和已发布反馈 |
| 考试与题库测试 | 验证题库权限、考试时间窗和试卷草稿 | 接口测试与边界用例 | 未授权题库不可用，未发布考试不可参加 |
| AI RAG 问答测试 | 验证课程上下文、引用和无来源处理 | Fake adapter 和后续向量检索测试 | 返回 citation 或 `NO_RELIABLE_SOURCE` |
| SSE 流式输出测试 | 验证 meta/delta/citation/done/error 顺序 | WebFlux 客户端、断流和取消测试 | 事件顺序正确，错误后关闭连接 |
| AI 引用来源测试 | 防止伪造或越权引用 | 使用未选课程资料、下线资料、版本变化资料测试 | 不返回未授权标题、片段和链接 |
| AI 人工确认测试 | 防止 AI 自动写正式数据 | 批改、摘要、预警、组卷页面验证 | 未确认时只保存草稿或建议 |
| 性能与并发基础测试 | 验证常用接口和 AI 限流 | 并发 50 用户基础压测，AI 限流测试 | 常用接口响应达标，AI 超限返回 429 |
| 演示数据测试 | 保证答辩演示可重复 | 固定账号、课程、作业、成绩、预警和 AI 示例 | 一键初始化后流程可复现 |

# 第十章 小组协作与开发计划

项目采用前端、后端业务、后端 AI、测试/部署协同推进。当前后端协作材料明确为两名后端成员：成员 A 负责 Biz 主链和正式业务事实，成员 B 负责 AI、考试题库、Gateway、Docker、CI 和联调。不要编造额外后端成员；如果后续实际成员名单确定，可在不改变职责边界的情况下替换角色名称。

Git 集成分支采用 `dev`，`backend-1`、`backend-2` 只作为个人备份或临时分支。功能分支从最新 `dev` 创建，经 PR 合入 `dev`。Flyway 历史迁移 V1 到 V3 不可修改，后续迁移使用 `VyyyyMMddHHmmss__description.sql`。接口契约、数据库设计、AI SSE 事件、内部 context 和错误码需要先评审再开发。

表 10-1 小组协作分工表

| 角色 | 主要职责 | 负责范围 | 协作规则 | 交付物 |
|---|---|---|---|---|
| 前端负责人 | Vue 3 页面、路由、组件、接口联调和状态管理 | 学生端、教师端、管理员端、Axios、Pinia、Element Plus 主题 | 只依赖已评审 API 契约；前端权限只做体验优化 | 页面实现、联调记录、前端构建 |
| 后端成员 A | Biz 主链和正式业务事实 | auth、course、assignment、grade、warning、forum、Biz shared、生产迁移 | 数据库、权限、Flyway、Biz shared 必须主审；不直接改 AI provider | 作业成绩闭环、预警、论坛、演示数据 |
| 后端成员 B | AI、考试题库、网关、部署联调 | edu-ai-service、edu-gateway、exam/question、Docker、CI、OpenAPI 聚合 | 不直接改作业、成绩、预警正式表；AI 不连 Biz MySQL | AI fake/真实 adapter、SSE、组卷建议、网关部署 |
| 测试/部署协同 | 自动测试、手工验收、演示环境和文档 | Maven verify、接口测试、权限测试、Docker Compose、启动说明 | 测试用例覆盖正常、异常、越权、AI 降级 | 测试报告、验收脚本、部署说明 |
| 文档维护 | 需求、设计、接口、数据库、答辩材料 | README、docs、OpenAPI、项目设计文档 | 文档与实现差异要标明“已实现/设计规划” | 设计文档、接口文档、答辩说明 |

共享模块修改规则：`edu-common` 只放技术协议，不放业务实体；父 POM、公共错误码、JWT claims、Gateway 安全、application 配置、Flyway 和 OpenAPI 需要单独 PR 说明影响。跨模块功能按“契约 -> 提供方 -> 消费方 -> 端到端验证”的顺序合并。

# 附录 A 自查与边界说明

本文全文使用“在线教育辅助教学系统”和“第一组”，未使用其他项目名称，也未保留模板占位符。本文包含 6 个 Mermaid 图和 10 张编号表格，图表均为可编辑文本或 Word 表格，不使用截图。AI 相关设计均遵守“AI 生成建议、草稿、解释、引用和任务状态；正式业务结果由人工确认后写入”的边界。

真实项目文件已经确认：组号来自既有需求分析文件，当前后端已实现认证与课程学习基础模块，实际表包括认证、课程、课程教师、选课、章节、课时、资料、学习记录和课程审核。作业、考试、成绩、论坛、公告、学习预警和 AI 会话/任务表属于本设计文档基于老师选题、需求分析和 MVP 计划补足的后续设计范围。
"""


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "请在 Word 中按 Ctrl+A 后按 F9 更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    paragraph.add_run(" 页")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def split_table_row(line):
    return [c.strip().replace("<br>", "\n") for c in line.strip().strip("|").split("|")]


def add_markdown_table(doc, lines):
    rows = [split_table_row(line) for line in lines if not re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line)]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        if r_idx == 0:
            repeat_table_header(table.rows[r_idx])
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for j, part in enumerate(text.split("\n")):
                if j:
                    para.add_run().add_break()
                run = para.add_run(part)
                set_run_font(run, size=8.0 if col_count >= 6 else 9.0, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "E8EEF5")
    doc.add_paragraph()


def make_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", 1)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.2)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

    if "CaptionText" not in styles:
        style = styles.add_style("CaptionText", 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F4D78")
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("在线教育辅助教学系统")
    set_run_font(run, size=24, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目设计文档")
    set_run_font(run, size=22, bold=True, color="2E74B5")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("概要设计与详细设计结合版")
    set_run_font(run, size=13, color="1F4D78")
    doc.add_paragraph()

    meta = [
        ("项目名称", "在线教育辅助教学系统"),
        ("项目选题", "选题 8"),
        ("小组名称/组号", "第一组"),
        ("文档版本", "V1.0"),
        ("生成日期", "2026-07-08"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    repeat_table_header(table.rows[0])
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for c in range(2):
            cell = table.cell(i, c)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c == 0:
                shade_cell(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(c == 0))
    doc.add_paragraph()
    p = doc.add_paragraph("本文档承接需求分析说明书，用于指导系统架构、数据库、接口、前后端、AI 服务、测试联调和答辩说明。")
    set_paragraph_spacing(p, after=8, line=1.1)
    doc.add_page_break()


def render_markdown_to_docx(doc, md):
    lines = md.strip().splitlines()
    in_code = False
    code_lang = ""
    table_lines = []
    first_h1 = True

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = line.strip("`").strip()
            else:
                in_code = False
                code_lang = ""
                doc.add_paragraph()
            continue
        if in_code:
            p = doc.add_paragraph(line, style="CodeBlock")
            for r in p.runs:
                set_run_font(r, "Consolas", "Microsoft YaHei", size=8.5)
            continue
        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 2")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            continue
        if re.match(r"^[图表]\s+\d+-\d+", line):
            p = doc.add_paragraph(line.strip(), style="CaptionText")
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s+", "", line)
            run = p.add_run(text)
            set_run_font(run)
            set_paragraph_spacing(p, after=4, line=1.1)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_run_font(run)
            set_paragraph_spacing(p, after=4, line=1.1)
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run)
        set_paragraph_spacing(p, after=6, line=1.1)
    flush_table()


def build_docx():
    doc = Document()
    make_styles(doc)
    add_cover(doc)
    toc_heading = doc.add_paragraph("目录", style="Heading 1")
    p = doc.add_paragraph()
    add_toc(p)
    note = doc.add_paragraph("说明：本目录为 Word 目录字段，打开文档后按 Ctrl+A，再按 F9 可刷新页码与层级。")
    set_paragraph_spacing(note, after=6)
    doc.add_page_break()
    render_markdown_to_docx(doc, BODY_MD)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "在线教育辅助教学系统 项目设计文档 | 第一组"
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in footer.runs:
        set_run_font(run, size=9, color="64748B")
    add_page_number(doc.sections[0].footer.add_paragraph())
    doc.core_properties.title = "在线教育辅助教学系统_项目设计文档_第一组"
    doc.core_properties.subject = "项目设计文档"
    doc.core_properties.author = "第一组"
    doc.core_properties.keywords = "在线教育辅助教学系统, 项目设计文档, 第一组"
    doc.save(OUT_DOCX)


def write_markdown():
    toc_note = "\n\n## 目录\n\n在 Word 版本中已插入可更新目录字段。Markdown 版本可使用编辑器的大纲功能查看标题层级。\n\n"
    OUT_MD.write_text(COVER_MD.strip() + toc_note + BODY_MD.strip() + "\n", encoding="utf-8")


def main():
    write_markdown()
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
